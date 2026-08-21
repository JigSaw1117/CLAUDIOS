"""
Corrige el Informe Tecnico Consolidado del Taller 1.

Aplica sobre INFORME_TALLER_1_CLAUDIOS.docx las correcciones detectadas al
contrastarlo con el enunciado, conservando sus capturas de pantalla:

  1. URL del aplicativo desplegado (Netlify) y nota sobre la tecnologia elegida.
  2. Caso 1 Fase A - se anaden outliers y multicolinealidad (correlacion + VIF).
  3. Caso 1 Fase B - se documenta la Regresion Lineal Multiple, que faltaba.
  4. Caso 2 Fase A - se anaden outliers y multicolinealidad.
  5. Caso 2 Fase B - se anade la validacion cruzada a la comparativa.
  6. Caso 2 Fase C - el indice real es index.html, no wine_predictor_v2.html.
  7. Seccion 5 - descripcion actualizada del aplicativo de pagina unica.
  8. Seccion 5.1 - estructura del repositorio al dia.
  9. Seccion 7 - comparativa global corregida.

Ejecutar:  python corregir_informe.py
"""

import copy
import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

ARCHIVO = "INFORME_TALLER_1_CLAUDIOS.docx"
AZUL = RGBColor(0x1F, 0x4E, 0x79)
doc = Document(ARCHIVO)
cambios = []


# --------------------------------------------------------------- utilidades
def buscar(fragmento):
    """Devuelve el primer parrafo que contiene el fragmento."""
    for p in doc.paragraphs:
        if fragmento in p.text:
            return p
    raise LookupError(f"No encontrado: {fragmento[:60]!r}")


def reescribir(parrafo, texto):
    """Sustituye el texto del parrafo por completo.

    Los hipervinculos de Word son elementos <w:hyperlink>, no runs: hay que
    eliminarlos aparte o la URL antigua sobrevive al reemplazo.
    """
    for hl in parrafo._p.findall(qn("w:hyperlink")):
        parrafo._p.remove(hl)
    for r in parrafo.runs[1:]:
        r._element.getparent().remove(r._element)
    if parrafo.runs:
        parrafo.runs[0].text = texto
    else:
        parrafo.add_run(texto)


def enlace(parrafo, texto, url):
    """Anade un hipervinculo real al final del parrafo."""
    rid = parrafo.part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/"
             "relationships/hyperlink", is_external=True)
    hl = OxmlElement("w:hyperlink"); hl.set(qn("r:id"), rid)
    run = OxmlElement("w:r"); rPr = OxmlElement("w:rPr")
    c = OxmlElement("w:color"); c.set(qn("w:val"), "0563C1"); rPr.append(c)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    run.append(rPr)
    t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve"); t.text = texto
    run.append(t); hl.append(run); parrafo._p.append(hl)


def parrafo_tras(ancla, texto, negrita=None, estilo=None):
    """Inserta un parrafo nuevo justo despues del ancla.

    'negrita' es el texto literal que debe ir resaltado al principio; se calcula
    su longitud, en lugar de fijar un numero de caracteres a mano (que partia
    palabras por la mitad).
    """
    nuevo = doc.add_paragraph(style=estilo)
    if negrita:
        assert texto.startswith(negrita), f"prefijo no coincide: {negrita!r}"
        r = nuevo.add_run(negrita); r.bold = True
        nuevo.add_run(texto[len(negrita):])
    else:
        nuevo.add_run(texto)
    nuevo.paragraph_format.space_after = Pt(6)
    nuevo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    ancla._p.addnext(nuevo._p)
    return nuevo


def sombrear(celda, color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color)
    celda._tc.get_or_add_tcPr().append(shd)


def tabla_tras(ancla, filas):
    """Crea una tabla con cabecera sombreada y la mueve tras el ancla."""
    t = doc.add_table(rows=0, cols=len(filas[0]))
    t.style = "Table Grid"
    for i, fila in enumerate(filas):
        celdas = t.add_row().cells
        for j, txt in enumerate(fila):
            celdas[j].text = ""
            p = celdas[j].paragraphs[0]
            r = p.add_run(str(txt))
            r.font.size = Pt(9)
            if i == 0:
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                sombrear(celdas[j], "1F4E79")
            elif i % 2 == 0:
                sombrear(celdas[j], "F2F6FA")
            if j > 0:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    ancla._p.addnext(t._tbl)
    return t


def borrar_tabla_tras(ancla):
    """Elimina la primera tabla que sigue al parrafo ancla."""
    el = ancla._p.getnext()
    while el is not None and not el.tag.endswith("}tbl"):
        el = el.getnext()
    if el is not None:
        el.getparent().remove(el)
        return True
    return False


# ============================================================ 1. URL y Fase C
p = buscar("Aplicativo web desplegado (URL pública)")
reescribir(p, "Aplicativo web desplegado (URL pública): ")
enlace(p, "https://californication1.netlify.app/", "https://californication1.netlify.app/")
cambios.append("URL del aplicativo -> Netlify")

p = buscar("se publica directamente con GitHub Pages")
reescribir(p,
    "El aplicativo es un sitio 100 % estático (HTML + JavaScript, sin backend): los modelos "
    "entrenados en Python se exportan a archivos JSON/JS y el navegador reproduce el pipeline "
    "de preprocesamiento y predicción. Por ello se despliega en Netlify como sitio estático, "
    "sin proceso de compilación. Las capturas de la sección 6 corresponden al aplicativo en "
    "funcionamiento.")
nota = parrafo_tras(p,
    "Nota sobre la tecnología empleada. El enunciado sugiere Streamlit, Gradio o Flask/FastAPI. "
    "Se optó por HTML + JavaScript puro tras reimplementar el pipeline de inferencia en el "
    "cliente: el aplicativo no requiere servidor de Python, arranca de forma instantánea, no "
    "depende de un servicio que pueda hibernar y puede desplegarse en cualquier hosting "
    "estático. La lógica de inferencia está aislada en funciones propias, de modo que un port "
    "a Streamlit reutilizaría directamente los mismos archivos de modelo.",
    negrita="Nota sobre la tecnología empleada.")
cambios.append("Nota justificando la tecnologia de la Fase C")


# ================================================== 2. Caso 1 - Fase A ampliada
p = buscar("Partición: 80 % entrenamiento / 20 % prueba (16 512 / 4 128 filas)")

t = tabla_tras(p, [
    ["Variable", "VIF", "Diagnóstico"],
    ["households", "28.28", "Severa"],
    ["total_bedrooms", "26.88", "Severa"],
    ["total_rooms", "12.13", "Severa"],
    ["latitude", "8.83", "Moderada"],
    ["longitude", "8.70", "Moderada"],
    ["population", "6.26", "Moderada"],
    ["median_income", "1.69", "Aceptable"],
    ["housing_median_age", "1.26", "Aceptable"],
])
parrafo_tras(p,
    "Multicolinealidad: se calculó la matriz de correlación de Pearson y el Factor de Inflación "
    "de la Varianza (VIF_j = 1 / (1 − R²_j)). Se detectan correlaciones muy altas entre las "
    "variables de tamaño del distrito: total_bedrooms–households (r = +0.980), "
    "total_rooms–total_bedrooms (r = +0.930) y population–households (r = +0.907), además de "
    "longitude–latitude (r = −0.925) por la geografía diagonal de California.",
    negrita="Multicolinealidad:")
parrafo_tras(p,
    "Outliers: detectados por el método IQR (k = 1.5). Afectan sobre todo a las variables de "
    "tamaño del distrito: total_rooms 1 287 (6.24 %), total_bedrooms 1 271 (6.16 %), "
    "households 1 220 (5.91 %), population 1 196 (5.79 %) y median_income 681 (3.30 %). "
    "Decisión: se conservan, ya que corresponden a distritos genuinamente grandes o de rentas "
    "altas, no a errores de medición; eliminarlos sesgaría el modelo hacia el distrito promedio. "
    "Adicionalmente, 965 registros (4.68 %) presentan el objetivo censurado en US$ 500 001, un "
    "tope administrativo del censo que constituye una limitación conocida del dataset.",
    negrita="Outliers:")
cambios.append("Caso 1 Fase A: outliers + correlacion + tabla VIF")

# El parrafo interpretativo va despues de la tabla VIF recien insertada.
ancla_vif = doc.paragraphs[[i for i, q in enumerate(doc.paragraphs)
                            if "Multicolinealidad: se calculó" in q.text][0]]
parrafo_tras(ancla_vif,
    "Las cuatro variables de mayor VIF (households, total_bedrooms, total_rooms y population) "
    "miden todas el tamaño del distrito, por lo que la multicolinealidad es estructural y "
    "esperable. No impide predecir, pero desaconseja interpretar esos coeficientes de forma "
    "individual.")


# ================================================ 3. Caso 1 - Fase B reescrita
p = buscar("Sobre las variables numéricas estandarizadas se aplica una expansión")
reescribir(p,
    "Se entrenan las dos metodologías exigidas compartiendo exactamente el mismo "
    "preprocesamiento (imputación por mediana → estandarización → codificación one-hot) y la "
    "misma partición 80/20 con random_state=42. La única diferencia entre ambas es la expansión "
    "polinómica, de modo que cualquier diferencia de rendimiento es atribuible al modelo y no "
    "al tratamiento de los datos.")

p2 = parrafo_tras(p,
    "Regresión Lineal Múltiple (modelo base). Ajusta ŷ = β₀ + Σ βᵢzᵢ sobre las 8 variables "
    "numéricas estandarizadas más las 5 dummies de ocean_proximity, resultando 13 términos más "
    "intercepto, estimados por mínimos cuadrados ordinarios.",
    negrita="Regresión Lineal Múltiple (modelo base).")
p3 = parrafo_tras(p2,
    "Regresión Polinómica de grado 2. Aplica PolynomialFeatures(degree=2) sobre las variables "
    "numéricas ya estandarizadas, generando términos cuadráticos zᵢ² e interacciones zᵢzⱼ que, "
    "junto con las dummies, producen 49 términos más intercepto. Sigue siendo lineal en los "
    "parámetros, que es lo que define el método. El pipeline completo se implementa en "
    "entrenar.py y se serializa en modelo_casas.pkl y modelo.json.",
    negrita="Regresión Polinómica de grado 2.")

# Sustituir la tabla antigua por una comparativa completa.
borrar_tabla_tras(p3)
tabla_tras(p3, [
    ["Métrica (prueba)", "Lineal múltiple", "Polinómica grado 2"],
    ["Términos", "13", "49"],
    ["R² entrenamiento", "0.6497", "0.7053"],
    ["R² prueba", "0.6254", "0.6570"],
    ["R² validación cruzada 5-fold", "0.6477 ± 0.013", "0.6262 ± 0.133"],
    ["Brecha entrenamiento − CV", "0.0020", "0.0791"],
    ["RMSE", "US$ 70 059.19", "US$ 67 043.55"],
    ["MAE", "US$ 50 670.49", "US$ 46 981.22"],
])
cambios.append("Caso 1 Fase B: se documenta la lineal multiple + tabla comparativa")

p = buscar("La expansión polinómica de grado 2 mejora el poder predictivo")
reescribir(p,
    "En el conjunto de prueba la expansión polinómica mejora el ajuste: +0.0316 de R² "
    "(+3.16 puntos porcentuales) y una reducción de US$ 3 689 en el MAE. Con 20 640 registros "
    "hay muestra suficiente para estimar 49 coeficientes, y las interacciones entre coordenadas "
    "geográficas e ingreso capturan la estructura espacial no lineal del precio de la vivienda. "
    "El error relativo del modelo polinómico es del 22.7 % sobre el precio medio.")
parrafo_tras(p,
    "Sin embargo, la lectura no es unánime y conviene reportarla completa: en validación cruzada "
    "5-fold el orden se invierte (0.6477 el lineal frente a 0.6262 el polinómico) y, sobre todo, "
    "la desviación entre particiones del polinomio es un orden de magnitud mayor (±0.133 frente "
    "a ±0.013), con una brecha entrenamiento−CV que pasa de 0.0020 a 0.0791. Es decir: el "
    "polinomio es más potente pero considerablemente menos estable. Se adopta como modelo "
    "desplegado por su mejor error de prueba, dejando constancia de esa inestabilidad. El "
    "repositorio incluye además la carpeta CASO_1/1_REGRESION_MULTIPLE con una variante "
    "avanzada de la regresión lineal múltiple (ingeniería de descriptores y regularización "
    "Ridge) que alcanza un R² de prueba de 0.6939.")
cambios.append("Caso 1 Fase B: matiz de validacion cruzada e inestabilidad")


# ================================================== 4. Caso 2 - Fase A ampliada
p = buscar("Partición: 80 % entrenamiento / 20 % prueba (914 / 229 muestras)")
parrafo_tras(p,
    "Multicolinealidad: ningún VIF supera el umbral de 10, por lo que este dataset no presenta "
    "multicolinealidad severa, a diferencia de los casos 1 y 3. Los valores más altos son "
    "fixed acidity (7.78) y density (6.60), ambos en rango moderado. Las correlaciones más "
    "fuertes son fixed acidity–pH (r = −0.685), fixed acidity–density (r = +0.682) y "
    "fixed acidity–citric acid (r = +0.673). Respecto al objetivo, las propiedades más "
    "asociadas a la calidad son alcohol (r = +0.485) y volatile acidity (r = −0.407).",
    negrita="Multicolinealidad:")
parrafo_tras(p,
    "Outliers: detectados por el método IQR (k = 1.5); los más numerosos aparecen en "
    "total sulfur dioxide, chlorides, sulphates y residual sugar. Decisión: se conservan, por "
    "tratarse de composiciones químicas reales —un vino puede presentar acidez volátil elevada "
    "y seguir siendo un vino válido— y porque con solo 1 143 muestras su eliminación reduciría "
    "una muestra ya pequeña y sesgaría el modelo hacia el vino promedio.",
    negrita="Outliers:")
cambios.append("Caso 2 Fase A: outliers + multicolinealidad")


# ================================================ 5. Caso 2 - Fase B ampliada
p = buscar("El resultado ilustra un caso clásico de sobreajuste")
reescribir(p,
    "El resultado ilustra un caso clásico de sobreajuste: la expansión polinómica multiplica "
    "los términos por siete (11 → 77) y mejora el ajuste en entrenamiento (+7.69 puntos de R²), "
    "pero empeora en prueba (−3.62 puntos) y se desploma en validación cruzada, que cae de "
    "0.3462 ± 0.043 a 0.2316 ± 0.147. La brecha entrenamiento−CV se multiplica por seis "
    "(0.0360 → 0.2275). Con 914 muestras de entrenamiento, 77 parámetros son sencillamente "
    "demasiados: el modelo memoriza el ruido de la muestra en lugar de aprender la relación.")
parrafo_tras(p,
    "Modelo recomendado: la Regresión Lineal Múltiple. El techo de R² ≈ 0.32 no es un defecto "
    "del modelo sino del problema: la calidad es una puntuación subjetiva y discreta asignada "
    "por catadores —concentrada en los valores 5 y 6— y 11 propiedades químicas no pueden "
    "explicar por completo una valoración sensorial.",
    negrita="Modelo recomendado: la Regresión Lineal Múltiple.")
cambios.append("Caso 2 Fase B: validacion cruzada y recomendacion explicita")

p = buscar("wine_predictor_v2.html, servido como índice del caso")
reescribir(p,
    "El aplicativo «Wine Quality Predictor» (CASO_2/index.html) es un comparador interactivo con "
    "tres modos: Solo Lineal, Comparar Ambos y Solo Polinómico. Permite ajustar las 11 "
    "propiedades con campos y deslizadores acotados al rango real del dataset, muestra la "
    "predicción de cada modelo, la fórmula con los coeficientes y las métricas de ambos. Los dos "
    "modelos se cargan mediante fetch desde wine_model_linear.json y wine_model_poly.json: "
    "ningún coeficiente está incrustado en el HTML.")
cambios.append("Caso 2 Fase C: indice real y carga externa de modelos")


# ============================================ 6. Seccion 5 - aplicativo actual
p = buscar("una portada institucional (index.html) presenta la matriz de casos")
reescribir(p,
    "La Fase C se resuelve con una aplicación de página única: index.html presenta una barra "
    "superior con botones que alternan entre las siete vistas del taller (Inicio, Caso 1 "
    "múltiple, Caso 1 polinómica, Caso 2, Caso 3 resumen, Caso 3 múltiple y Caso 3 polinomial) "
    "sin recargar ni abandonar la página. La portada institucional se conserva como portada.html "
    "y es la vista de inicio. El usuario puede así seleccionar cualquiera de los tres casos de "
    "estudio e ingresar parámetros para obtener predicciones en tiempo real.")

p = buscar("Los modelos entrenados en Python se exportan con sus medias")
reescribir(p,
    "Arquitectura: HTML + JavaScript puro, sin dependencias ni proceso de build. Los modelos "
    "entrenados en Python se exportan con sus medias, desviaciones, exponentes de la expansión "
    "polinómica, coeficientes e intercepto a archivos externos (modelo.json, "
    "wine_model_linear.json, wine_model_poly.json y los model.js de cada metodología), y el "
    "navegador reproduce exactamente el pipeline de preprocesamiento y predicción. Ninguna "
    "página lleva el modelo incrustado: todas lo consumen desde su archivo, mediante fetch o "
    "mediante <script src>. Por ello el sitio debe servirse por HTTP y no abrirse con doble clic.")
cambios.append("Seccion 5: aplicativo de pagina unica y consumo externo del modelo")


# ============================================= 7. Seccion 5.1 - estructura
# El arbol ocupa varios parrafos seguidos: se localiza el primero de forma
# exacta (no por 'CLAUDIOS/', que tambien aparece dentro de las URLs) y se
# eliminan los restantes tras reescribirlo.
idx = next(i for i, q in enumerate(doc.paragraphs) if q.text.strip() == "CLAUDIOS/")
p = doc.paragraphs[idx]
for extra in [q for q in doc.paragraphs[idx + 1: idx + 16]
              if q.text.strip().startswith(("├──", "└──", "│", "    ├", "    └"))]:
    extra._p.getparent().remove(extra._p)
reescribir(p,
    "CLAUDIOS/\n"
    "├── index.html                     · aplicativo de página única con botones por caso\n"
    "├── portada.html                   · portada institucional (vista de inicio)\n"
    "├── README.md · requirements.txt · netlify.toml\n"
    "├── CASO_1/                        · California Housing\n"
    "│   ├── CASO_1_analisis.ipynb      · cuaderno: Fases A y B\n"
    "│   ├── entrenar.py · analisis_fases_a_b.py · interfaz.py\n"
    "│   ├── index.html · modelo.json · modelo_casas.pkl · housing.csv\n"
    "│   └── 1_REGRESION_MULTIPLE/      · train.py · model.js · index.html · figuras/\n"
    "├── CASO_2/                        · Wine Quality\n"
    "│   ├── CASO_2_analisis.ipynb      · cuaderno: Fases A y B\n"
    "│   ├── train_wine_model.py · analisis_fases_a_b.py\n"
    "│   └── index.html · wine_model_linear.json · wine_model_poly.json · WineQT.csv\n"
    "└── CASO_3/                        · Diabetes\n"
    "    ├── CASO_3_analisis.ipynb      · cuaderno: Fases A y B\n"
    "    ├── index.html · INFORME_CONSOLIDADO.pdf · PLAN.md\n"
    "    ├── 1_REGRESION_MULTIPLE/      · train.py · model.js · index.html · figuras/\n"
    "    └── 2_REGRESION_POLINOMIAL/    · train.py · model.js · index.html · figuras/")
cambios.append("Seccion 5.1: estructura del repositorio actualizada")


# ======================================== 8. Seccion 7 - comparativa global
for t in doc.tables:
    if t.rows and "Aspecto" in t.rows[0].cells[0].text:
        for fila in t.rows:
            et = fila.cells[0].text.strip()
            if et.startswith("¿Aporta el polinomio?"):
                fila.cells[1].text = "En prueba sí (+3.16 p.p.); en CV no"
                fila.cells[2].text = "No (sobreajuste: −3.6 p.p.)"
                fila.cells[3].text = "No (indistinguible)"
            elif et.startswith("Mejor modelo"):
                fila.cells[1].text = "Polinómico gr. 2 (por error de prueba)"
        cambios.append("Seccion 7: fila del polinomio del Caso 1 matizada")
        break

p = buscar("La comparación entre dominios deja tres lecciones generales")
reescribir(p,
    "La comparación entre dominios deja tres lecciones generales. Primera: la utilidad de la "
    "regresión polinomial depende del tamaño muestral y de la estructura real de los datos. Con "
    "20 640 registros y relaciones espaciales genuinamente no lineales (Caso 1) el grado 2 "
    "mejora el error de prueba, aunque a costa de una notable pérdida de estabilidad entre "
    "particiones; con 1 143 muestras (Caso 2) la expansión a 77 términos ya sobreajusta; y con "
    "442 pacientes y multicolinealidad estructural (Caso 3) la expansión no encuentra nada que "
    "capturar.")
cambios.append("Seccion 7: matiz sobre la estabilidad del polinomio")


doc.save(ARCHIVO)
print(f"{ARCHIVO} actualizado ({os.path.getsize(ARCHIVO)/1024/1024:.2f} MB)\n")
for i, c in enumerate(cambios, 1):
    print(f"  {i}. {c}")
