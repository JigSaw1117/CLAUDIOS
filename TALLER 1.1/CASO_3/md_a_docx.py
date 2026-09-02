"""
Conversor de los informes Markdown a documentos Word (.docx) con formato profesional.

Genera carátula, tabla de contenido, encabezados de estilo Word, tablas con cabecera
sombreada, bloques de código monoespaciados, citas, listas, figuras incrustadas y
numeración de páginas en el pie.

Ejecutar:  python md_a_docx.py
Requiere:  python-docx
"""

import os
import re
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ----------------------------------------------------------------- configuración
AZUL = RGBColor(0x1F, 0x4E, 0x79)
GRIS = RGBColor(0x44, 0x4444 % 256, 0x4C)
MONO = "Consolas"
ANCHO_UTIL_CM = 16.0          # A4 (21 cm) menos margenes de 2.5 cm
SOMBRA_CABECERA = "1F4E79"
SOMBRA_CODIGO = "F2F2F2"
SOMBRA_CITA = "FFF7E6"

DOCUMENTOS = [
    ("INFORME.md", "INFORME_CONSOLIDADO.docx",
     "Informe Técnico Consolidado",
     "Caso 3 — Diabetes · Regresión Lineal Múltiple y Regresión Polinomial"),
    ("1_REGRESION_MULTIPLE/INFORME_1_REGRESION_MULTIPLE.md",
     "1_REGRESION_MULTIPLE/INFORME_1_REGRESION_MULTIPLE.docx",
     "Informe 1 — Regresión Lineal Múltiple",
     "Caso 3 — Diabetes · Fases A, B y C"),
    ("2_REGRESION_POLINOMIAL/INFORME_2_REGRESION_POLINOMIAL.md",
     "2_REGRESION_POLINOMIAL/INFORME_2_REGRESION_POLINOMIAL.docx",
     "Informe 2 — Regresión Polinomial",
     "Caso 3 — Diabetes · Fases A, B y C"),
]

# Los enlaces entre informes deben apuntar al .docx cuando el destino es Word.
EQUIVALENTE_DOCX = {
    "INFORME.md": "../INFORME_CONSOLIDADO.docx",
    "INFORME_1_REGRESION_MULTIPLE.md":
        "1_REGRESION_MULTIPLE/INFORME_1_REGRESION_MULTIPLE.docx",
    "INFORME_2_REGRESION_POLINOMIAL.md":
        "2_REGRESION_POLINOMIAL/INFORME_2_REGRESION_POLINOMIAL.docx",
}

PIES_FIGURA = {
    "a3_boxplots.png": "Diagramas de caja por variable (método IQR)",
    "a3_target.png": "Distribución y gráfico Q-Q de la variable objetivo",
    "a4_correlacion.png": "Matriz de correlación de Pearson",
    "a4_vif.png": "Factor de Inflación de la Varianza (VIF)",
    "b4_curva_validacion.png": "Curva de validación por grado polinomial",
    "b_curva_validacion.png": "Curva de validación por grado polinomial",
    "b_brecha.png": "Brecha entre entrenamiento y validación por grado",
    "a_correlacion.png": "Matriz de correlación de Pearson",
    "b_residuos.png": "Diagnóstico de residuos del modelo lineal",
    "b_real_vs_predicho.png": "Valores reales frente a valores predichos",
    "b_coeficientes.png": "Peso de cada variable en el modelo",
    "b5_residuos.png": "Diagnóstico de residuos del modelo lineal",
    "b5_real_vs_predicho.png": "Valores reales frente a valores predichos",
    "b5_coeficientes.png": "Peso de cada variable en el modelo",
}


# ------------------------------------------------------------------- utilidades
def sombrear(elemento, color_hex):
    """Aplica sombreado de fondo a una celda o a un párrafo."""
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")       # nunca 'solid': se renderiza en negro
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    elemento.append(shd)


def borde_parrafo(parrafo, lado="bottom", tamano=6, color="1F4E79"):
    pPr = parrafo._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    b = OxmlElement(f"w:{lado}")
    b.set(qn("w:val"), "single")
    b.set(qn("w:sz"), str(tamano))
    b.set(qn("w:space"), "4")
    b.set(qn("w:color"), color)
    pbdr.append(b)
    pPr.append(pbdr)


def campo(parrafo, instruccion):
    """Inserta un campo de Word (TOC, PAGE, NUMPAGES)."""
    r1 = OxmlElement("w:r"); f1 = OxmlElement("w:fldChar")
    f1.set(qn("w:fldCharType"), "begin"); r1.append(f1)
    r2 = OxmlElement("w:r"); t = OxmlElement("w:instrText")
    t.set(qn("xml:space"), "preserve"); t.text = instruccion; r2.append(t)
    r3 = OxmlElement("w:r"); f3 = OxmlElement("w:fldChar")
    f3.set(qn("w:fldCharType"), "separate"); r3.append(f3)
    r4 = OxmlElement("w:r"); tt = OxmlElement("w:t"); tt.text = " "; r4.append(tt)
    r5 = OxmlElement("w:r"); f5 = OxmlElement("w:fldChar")
    f5.set(qn("w:fldCharType"), "end"); r5.append(f5)
    for r in (r1, r2, r3, r4, r5):
        parrafo._p.append(r)


def hipervinculo(parrafo, texto, url):
    """python-docx no expone hipervínculos: se crea la relación a mano."""
    rid = parrafo.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), rid)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color"); color.set(qn("w:val"), "0563C1"); rPr.append(color)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    run.append(rPr)
    t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve"); t.text = texto
    run.append(t)
    link.append(run)
    parrafo._p.append(link)


# --------------------------------------------------------- parseo de línea inline
PATRON_INLINE = re.compile(
    r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`|\[[^\]]+\]\([^)]+\)|https?://\S+)")


def escribir_inline(parrafo, texto, base_negrita=False, base_tam=None,
                    base_cursiva=False):
    """Escribe texto con formato **negrita**, *cursiva*, `código` y enlaces.

    El parseo es recursivo: `codigo` y *cursiva* dentro de **negrita** se
    reconocen igualmente (Markdown permite anidarlos).
    """
    for parte in PATRON_INLINE.split(texto):
        if not parte:
            continue

        m = re.fullmatch(r"\[([^\]]+)\]\(([^)]+)\)", parte)
        if m:
            # Dentro del Word, los enlaces entre informes apuntan al .docx.
            hipervinculo(parrafo, m.group(1), EQUIVALENTE_DOCX.get(m.group(2), m.group(2)))
            continue
        if parte.startswith("http"):
            hipervinculo(parrafo, parte, parte)
            continue

        if parte.startswith("**") and parte.endswith("**"):
            escribir_inline(parrafo, parte[2:-2], True, base_tam, base_cursiva)
            continue
        if parte.startswith("*") and parte.endswith("*") and len(parte) > 2:
            escribir_inline(parrafo, parte[1:-1], base_negrita, base_tam, True)
            continue

        if parte.startswith("`") and parte.endswith("`"):
            run = parrafo.add_run(parte[1:-1])
            run.font.name = MONO
            run.font.size = Pt((base_tam or 10.5) - 1)
            run.font.color.rgb = RGBColor(0xA3, 0x15, 0x15)
        else:
            run = parrafo.add_run(parte)
            if base_tam:
                run.font.size = Pt(base_tam)
        run.bold = base_negrita
        run.italic = base_cursiva


def limpiar(texto):
    """Quita marcas de formato para medir el ancho de una celda."""
    return re.sub(r"[*`]", "", texto)


# ----------------------------------------------------------------- construcción
def png_tamano(ruta):
    """Lee ancho y alto de un PNG desde su cabecera IHDR."""
    with open(ruta, "rb") as f:
        cab = f.read(24)
    return int.from_bytes(cab[16:20], "big"), int.from_bytes(cab[20:24], "big")


def agregar_figura(doc, ruta, base):
    if not os.path.exists(ruta):
        return
    ancho_px, alto_px = png_tamano(ruta)
    ancho = min(ANCHO_UTIL_CM, 16.0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(ruta, width=Cm(ancho))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(f"Figura. {PIES_FIGURA.get(base, base)}")
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = GRIS
    cap.paragraph_format.space_after = Pt(14)


def agregar_tabla(doc, filas):
    encabezado, cuerpo = filas[0], filas[1:]
    n = len(encabezado)
    tabla = doc.add_table(rows=1, cols=n)
    tabla.style = "Table Grid"
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    tabla.autofit = False

    # Ancho por columna proporcional al contenido, pero garantizando que la
    # palabra más larga entre completa (si no, Word parte "Términos" en "Términ/s").
    CM_POR_CARACTER = 0.185          # aproximación para Calibri 9.5 pt
    CM_POR_CARACTER_NEGRITA = 0.215  # la cabecera va en negrita: ocupa más
    largos, minimos = [], []
    for j in range(n):
        cab = limpiar(encabezado[j])
        cuerpo_txt = [limpiar(f[j]) if j < len(f) else "" for f in cuerpo]
        largos.append(max([len(cab)] + [len(t) for t in cuerpo_txt]) or 1)
        pal_cab = max((len(p) for p in cab.split()), default=1)
        pal_cuerpo = max((len(p) for t in cuerpo_txt for p in t.split()), default=1)
        minimos.append(max(pal_cab * CM_POR_CARACTER_NEGRITA,
                           pal_cuerpo * CM_POR_CARACTER) + 0.30)

    total = sum(largos)
    anchos = [max(minimos[j], min(7.5, ANCHO_UTIL_CM * largos[j] / total))
              for j in range(n)]

    # Si los mínimos desbordan el ancho útil, se recorta solo lo que sobra de ellos.
    exceso = sum(anchos) - ANCHO_UTIL_CM
    if exceso > 0:
        holgura = [anchos[j] - minimos[j] for j in range(n)]
        total_holgura = sum(holgura)
        if total_holgura > 0:
            factor = max(0.0, 1 - exceso / total_holgura)
            anchos = [minimos[j] + holgura[j] * factor for j in range(n)]
        else:
            escala = ANCHO_UTIL_CM / sum(anchos)
            anchos = [a * escala for a in anchos]
    else:
        escala = ANCHO_UTIL_CM / sum(anchos)
        anchos = [a * escala for a in anchos]

    for j, texto in enumerate(encabezado):
        celda = tabla.rows[0].cells[j]
        celda.width = Cm(anchos[j])
        celda.text = ""
        p = celda.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        escribir_inline(p, texto, base_negrita=True, base_tam=9.5)
        for r in p.runs:
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        sombrear(celda._tc.get_or_add_tcPr(), SOMBRA_CABECERA)

    for i, fila in enumerate(cuerpo):
        celdas = tabla.add_row().cells
        for j in range(n):
            celda = celdas[j]
            celda.width = Cm(anchos[j])
            p = celda.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            escribir_inline(p, fila[j] if j < len(fila) else "", base_tam=9.5)
            if j > 0:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if i % 2 == 1:                        # filas alternas sombreadas
            for celda in celdas:
                sombrear(celda._tc.get_or_add_tcPr(), "F7F9FC")
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def agregar_codigo(doc, lineas, lenguaje):
    for k, linea in enumerate(lineas):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.left_indent = Cm(0.5)
        pf.space_before = Pt(6 if k == 0 else 0)
        pf.space_after = Pt(6 if k == len(lineas) - 1 else 0)
        pf.line_spacing = 1.0
        run = p.add_run(linea if linea else " ")
        run.font.name = MONO
        run.font.size = Pt(9)
        sombrear(p._p.get_or_add_pPr(), SOMBRA_CODIGO)


def caratula(doc, titulo, subtitulo):
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("REGRESIÓN LINEAL MÚLTIPLE Y POLINOMIAL")
    r.bold = True; r.font.size = Pt(12); r.font.color.rgb = GRIS

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(titulo)
    r.bold = True; r.font.size = Pt(26); r.font.color.rgb = AZUL
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitulo)
    r.font.size = Pt(13); r.font.color.rgb = GRIS
    borde_parrafo(p, "bottom", 12)
    p.paragraph_format.space_after = Pt(30)

    datos = [
        ("Dataset", "sklearn.datasets.load_diabetes (442 muestras, 10 variables)"),
        ("Contexto", "Progresión cuantitativa de la enfermedad un año después del inicio"),
        ("Metodología", "Regresión Lineal Múltiple y Regresión Polinomial"),
        ("Partición", "75 % entrenamiento / 25 % prueba (random_state = 42)"),
        ("Herramientas", "Python 3.13 · scikit-learn 1.9 · pandas · NumPy · SciPy · Matplotlib"),
        ("Reproducción", "python train.py"),
    ]
    t = doc.add_table(rows=0, cols=2)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for etiqueta, valor in datos:
        c = t.add_row().cells
        c[0].width = Cm(4.2); c[1].width = Cm(11.8)
        p0 = c[0].paragraphs[0]; r0 = p0.add_run(etiqueta)
        r0.bold = True; r0.font.size = Pt(10); r0.font.color.rgb = AZUL
        sombrear(c[0]._tc.get_or_add_tcPr(), "EAF1F8")
        p1 = c[1].paragraphs[0]
        escribir_inline(p1, valor, base_tam=10)

    doc.add_paragraph().paragraph_format.space_after = Pt(24)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("INTEGRANTES Y DIVISIÓN DE FUNCIONES")
    r.bold = True; r.font.size = Pt(11); r.font.color.rgb = AZUL
    p.paragraph_format.space_after = Pt(8)

    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for j, cab in enumerate(["Integrante", "Función desempeñada"]):
        c = t.rows[0].cells[j]
        c.width = Cm([7.0, 9.0][j])
        p = c.paragraphs[0]; r = p.add_run(cab)
        r.bold = True; r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        sombrear(c._tc.get_or_add_tcPr(), SOMBRA_CABECERA)
    for _ in range(4):
        c = t.add_row().cells
        c[0].width = Cm(7.0); c[1].width = Cm(9.0)
        for celda in c:
            celda.paragraphs[0].add_run(" ").font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("(completar antes de la entrega)")
    r.italic = True; r.font.size = Pt(8.5); r.font.color.rgb = GRIS

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def indice(doc):
    p = doc.add_paragraph()
    r = p.add_run("Tabla de contenido")
    r.bold = True; r.font.size = Pt(16); r.font.color.rgb = AZUL
    borde_parrafo(p, "bottom", 6)
    p.paragraph_format.space_after = Pt(10)

    p = doc.add_paragraph()
    campo(p, r'TOC \o "1-3" \h \z \u')

    p = doc.add_paragraph()
    r = p.add_run("Para actualizar el índice en Word: clic derecho sobre él → "
                  "Actualizar campos → Actualizar toda la tabla.")
    r.italic = True; r.font.size = Pt(8.5); r.font.color.rgb = GRIS

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def pie_de_pagina(doc, texto_izquierda):
    pie = doc.sections[0].footer
    p = pie.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(texto_izquierda + "   ·   Página ")
    r.font.size = Pt(8.5); r.font.color.rgb = GRIS
    campo(p, "PAGE")
    r = p.add_run(" de ")
    r.font.size = Pt(8.5); r.font.color.rgb = GRIS
    campo(p, "NUMPAGES")
    for run in p.runs:
        run.font.size = Pt(8.5)
        run.font.color.rgb = GRIS


# ------------------------------------------------------------------ conversión
def convertir(md, salida, titulo, subtitulo):
    base_dir = os.path.dirname(md)
    with open(md, encoding="utf-8") as f:
        lineas = f.read().split("\n")

    # El encabezado del Markdown (títulos H1 y bloque de metadatos) ya está en la
    # carátula: se descarta para no repetirlo. Solo se conserva la nota en cita.
    corte = next((k for k, l in enumerate(lineas) if re.fullmatch(r"-{3,}", l.strip())), 0)
    if corte:
        lineas = [l for l in lineas[:corte] if l.strip().startswith(">")] + lineas[corte + 1:]

    doc = Document()

    est = doc.styles["Normal"]
    est.font.name = "Calibri"
    est.font.size = Pt(10.5)
    est.paragraph_format.space_after = Pt(6)
    est.paragraph_format.line_spacing = 1.15

    for nivel, tam in ((1, 18), (2, 14), (3, 11.5)):
        s = doc.styles[f"Heading {nivel}"]
        s.font.name = "Calibri"
        s.font.size = Pt(tam)
        s.font.bold = True
        s.font.color.rgb = AZUL

    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21), Cm(29.7)          # A4
    sec.left_margin = sec.right_margin = Cm(2.5)
    sec.top_margin = sec.bottom_margin = Cm(2.2)

    caratula(doc, titulo, subtitulo)
    indice(doc)
    pie_de_pagina(doc, titulo)

    i = 0
    primer_h1 = True
    figuras_puestas = set()

    while i < len(lineas):
        linea = lineas[i]
        desnuda = linea.strip()

        # --- bloque de código
        if desnuda.startswith("```"):
            lenguaje = desnuda[3:].strip()
            i += 1
            buffer = []
            while i < len(lineas) and not lineas[i].strip().startswith("```"):
                buffer.append(lineas[i])
                i += 1
            agregar_codigo(doc, buffer, lenguaje)
            i += 1
            continue

        # --- tabla
        if desnuda.startswith("|") and i + 1 < len(lineas) \
                and re.match(r"^\|[\s:|-]+\|$", lineas[i + 1].strip()):
            filas = []
            while i < len(lineas) and lineas[i].strip().startswith("|"):
                cruda = lineas[i].strip()
                if not re.match(r"^\|[\s:|-]+\|$", cruda):
                    filas.append([c.strip() for c in cruda.strip("|").split("|")])
                i += 1
            agregar_tabla(doc, filas)
            continue

        # --- encabezados
        m = re.match(r"^(#{1,4})\s+(.*)$", desnuda)
        if m:
            nivel = min(len(m.group(1)), 3)
            texto = re.sub(r"[*`]", "", m.group(2))
            if nivel == 1 and not primer_h1:
                doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            primer_h1 = False
            h = doc.add_heading(level=nivel)
            r = h.add_run(texto)
            r.font.color.rgb = AZUL
            h.paragraph_format.space_before = Pt(14 if nivel == 1 else 10)
            h.paragraph_format.space_after = Pt(6)
            if nivel == 1:
                borde_parrafo(h, "bottom", 8)
            i += 1
            continue

        # --- separador horizontal: se omite (los H1 ya abren página)
        if re.fullmatch(r"-{3,}", desnuda):
            i += 1
            continue

        # --- cita
        if desnuda.startswith(">"):
            buffer = []
            while i < len(lineas) and lineas[i].strip().startswith(">"):
                buffer.append(lineas[i].strip().lstrip(">").strip())
                i += 1
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.left_indent = Cm(0.6); pf.right_indent = Cm(0.3)
            pf.space_before = Pt(8); pf.space_after = Pt(8)
            escribir_inline(p, " ".join(buffer))
            borde_parrafo(p, "left", 18, "FFB454")
            sombrear(p._p.get_or_add_pPr(), SOMBRA_CITA)
            continue

        # --- referencia a figura: se incrusta la imagen
        if re.match(r"^\*Figuras?:", desnuda):
            for ruta in re.findall(r"figuras/([\w.]+\.png)", desnuda):
                completa = os.path.join(base_dir, "figuras", ruta)
                if ruta not in figuras_puestas and os.path.exists(completa):
                    agregar_figura(doc, completa, ruta)
                    figuras_puestas.add(ruta)
            i += 1
            continue

        # --- listas
        m = re.match(r"^(\s*)([-*]|\d+\.)\s+(.*)$", linea)
        if m:
            sangria, marca, texto = m.group(1), m.group(2), m.group(3)
            nivel = len(sangria) // 2
            numerada = marca[0].isdigit()
            if numerada:
                # Numeración literal del Markdown: la automática de Word continúa la
                # cuenta de la lista anterior del documento y desordena los números.
                p = doc.add_paragraph()
                pf = p.paragraph_format
                pf.left_indent = Cm(1.0 + 0.5 * nivel)
                pf.first_line_indent = Cm(-0.6)
                p.add_run(marca + " ").bold = True
            else:
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.left_indent = Cm(0.75 + 0.5 * nivel)
            p.paragraph_format.space_after = Pt(3)
            # Las continuaciones de línea del mismo ítem se unen.
            i += 1
            while i < len(lineas) and lineas[i].startswith("   ") \
                    and lineas[i].strip() and not re.match(r"^\s*([-*]|\d+\.)\s", lineas[i]):
                texto += " " + lineas[i].strip()
                i += 1
            escribir_inline(p, texto)
            continue

        # --- línea en blanco
        if not desnuda:
            i += 1
            continue

        # --- párrafo (se unen las líneas hasta la siguiente en blanco)
        buffer = [desnuda]
        i += 1
        while i < len(lineas) and lineas[i].strip() \
                and not re.match(r"^(#{1,4}\s|\||>|```|\s*([-*]|\d+\.)\s|-{3,}$|\*Figuras?:)",
                                 lineas[i].strip()):
            buffer.append(lineas[i].strip())
            i += 1
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        escribir_inline(p, " ".join(buffer))

    doc.save(salida)
    return salida


if __name__ == "__main__":
    for md, salida, titulo, subtitulo in DOCUMENTOS:
        if not os.path.exists(md):
            print(f"  AVISO: no existe {md}, se omite")
            continue
        ruta = convertir(md, salida, titulo, subtitulo)
        print(f"  {ruta:42s} {os.path.getsize(ruta)/1024:7.1f} kB")
    print("\nDocumentos Word generados.")
