"""
Genera INFORME_POTABILIDAD_AGUA.docx a partir del contenido de
INFORME_POTABILIDAD_AGUA.md, con formato profesional: caratula, tabla de
contenido, tablas con cabecera, figuras incrustadas y numeracion de paginas.

Ejecutar:  python build_informe.py
Requiere:  python-docx
"""

import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

AZUL = RGBColor(0x0D, 0x3C, 0x6C)
GRIS = RGBColor(0x47, 0x55, 0x69)
ANCHO_UTIL_CM = 16.0
SOMBRA_CABECERA = "0D3C6C"

MD_PATH = "INFORME_POTABILIDAD_AGUA.md"
OUT_PATH = "INFORME_POTABILIDAD_AGUA.docx"


# ----------------------------------------------------------------- utilidades
def sombrear_celda(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def set_cell_width(cell, cm):
    cell.width = Cm(cm)


def add_page_number_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld1)
    run._r.append(instr)
    run._r.append(fld2)


def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-2" \\h \\z \\u'
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "separate")
    txt = OxmlElement("w:t")
    txt.text = "Actualice la tabla de contenido (clic derecho > Actualizar campo)."
    fld3 = OxmlElement("w:fldChar")
    fld3.set(qn("w:fldCharType"), "end")
    run._r.append(fld1)
    run._r.append(instr)
    run._r.append(fld2)
    run._r.append(txt)
    run._r.append(fld3)


def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


INLINE_TOKEN = re.compile(
    r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\)|\*[^*]+\*)"
)


def parse_inline(paragraph, text):
    """Negrita **t**, italica *t*, codigo `t` y enlaces [t](url) dentro de una linea."""
    for part in INLINE_TOKEN.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
        elif part.startswith("[") and "](" in part and part.endswith(")"):
            m = re.match(r"\[([^\]]+)\]\(([^)]+)\)", part)
            label, url = m.groups()
            label = label.strip("*`")
            add_hyperlink(paragraph, url, label)
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def add_table_from_md(doc, header, rows):
    table = doc.add_table(rows=1, cols=len(header))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    col_w = ANCHO_UTIL_CM / len(header)
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(header):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(h.strip())
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sombrear_celda(hdr_cells[i], SOMBRA_CABECERA)
        set_cell_width(hdr_cells[i], col_w)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            val = val.strip()
            plano = re.sub(r"[*`]", "", val)
            parse_inline(p, val)
            for run in p.runs:
                run.font.size = Pt(10)
            if re.match(r"^[+-]?[\d.,%]+$", plano):
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            set_cell_width(cells[i], col_w)
    doc.add_paragraph()


# ----------------------------------------------------------------- documento
doc = Document()

section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
add_page_number_footer(section)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

for lvl, size, color in [("Heading 1", 16, AZUL), ("Heading 2", 13, AZUL), ("Heading 3", 11.5, GRIS)]:
    hs = doc.styles[lvl]
    hs.font.name = "Calibri"
    hs.font.size = Pt(size)
    hs.font.color.rgb = color
    hs.font.bold = True

with open(MD_PATH, encoding="utf-8") as f:
    lines = f.read().splitlines()

# ------------------------------------------------------------------ caratula
i = 0
caratula_lines = []
while i < len(lines) and lines[i].strip() != "---":
    caratula_lines.append(lines[i])
    i += 1
i += 1  # salta el primer '---'

titulo_principal = caratula_lines[0].lstrip("# ").strip()
for _ in range(4):
    doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(titulo_principal)
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = AZUL
doc.add_paragraph()

for extra in caratula_lines[1:]:
    extra = extra.strip()
    if not extra:
        continue
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if extra.startswith("**Integrantes:**"):
        run = p.add_run("Integrantes:")
        run.bold = True
        continue
    if extra.startswith("- "):
        run = p.add_run(extra[2:])
        continue
    parse_inline(p, extra)

for _ in range(6):
    doc.add_paragraph()

doc.add_page_break()

toc_title = doc.add_heading("Tabla de Contenido", level=1)
add_toc(doc)
doc.add_page_break()

# ------------------------------------------------------------------- cuerpo
table_header = None
table_rows = []
in_table = False

while i < len(lines):
    line = lines[i]
    stripped = line.strip()

    if stripped.startswith("|"):
        cells = [c.strip() for c in stripped.strip("|").split("|")]
        if not in_table:
            table_header = cells
            in_table = True
        elif re.match(r"^:?-+:?$", cells[0].replace(" ", "")):
            pass  # fila separadora markdown, se ignora
        else:
            table_rows.append(cells)
        i += 1
        continue
    elif in_table:
        add_table_from_md(doc, table_header, table_rows)
        table_header, table_rows, in_table = None, [], False

    if stripped.startswith("```"):
        i += 1
        code_lines = []
        while i < len(lines) and not lines[i].strip().startswith("```"):
            code_lines.append(lines[i])
            i += 1
        p = doc.add_paragraph()
        run = p.add_run("\n".join(code_lines))
        run.font.name = "Consolas"
        run.font.size = Pt(9.5)
        p.paragraph_format.left_indent = Cm(0.5)
        i += 1
        continue

    m_img = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", stripped)
    if m_img:
        alt, path = m_img.groups()
        try:
            doc.add_picture(path, width=Cm(14.5))
            last_p = doc.paragraphs[-1]
            last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cap.add_run(alt)
            run.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = GRIS
        except Exception as e:
            doc.add_paragraph(f"[Figura no encontrada: {path}]")
        i += 1
        continue

    if stripped.startswith("#### "):
        doc.add_heading(stripped[5:], level=3)
    elif stripped.startswith("### "):
        doc.add_heading(stripped[4:], level=3)
    elif stripped.startswith("## "):
        doc.add_heading(stripped[3:], level=2)
    elif stripped.startswith("# "):
        doc.add_heading(stripped[2:], level=1)
    elif stripped.startswith("- "):
        p = doc.add_paragraph(style="List Bullet")
        parse_inline(p, stripped[2:])
    elif re.match(r"^\d+\.\s", stripped):
        numero = re.match(r"^(\d+)\.\s", stripped).group(1)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.65)
        p.paragraph_format.first_line_indent = Cm(-0.4)
        parse_inline(p, f"{numero}. " + re.sub(r"^\d+\.\s", "", stripped))
    elif stripped == "":
        pass
    else:
        p = doc.add_paragraph()
        parse_inline(p, stripped)

    i += 1

if in_table:
    add_table_from_md(doc, table_header, table_rows)

doc.save(OUT_PATH)
print(f"Guardado: {OUT_PATH}")
